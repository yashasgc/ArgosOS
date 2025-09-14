"""
IngestAgent - Handles file ingestion, text extraction, and AI processing
"""
import logging
import time
import json
from pathlib import Path
from typing import List, Optional, Tuple
from io import BytesIO

from sqlalchemy.orm import Session

from app.db.models import Document
from app.db.crud import DocumentCRUD, TagCRUD
from app.db.schemas import DocumentCreate
from app.llm.provider import LLMProvider
from app.utils.hash import compute_bytes_hash
from app.constants import (
    MAX_FILE_SIZE, MAX_TEXT_PREVIEW, MAX_CONTENT_PREVIEW, MAX_SUMMARY_PREVIEW,
    TEXT_ENCODINGS
)

logger = logging.getLogger(__name__)

# Import OCR and text extraction libraries
try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import PyMuPDF as fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class IngestAgent:
    """
    Agent responsible for ingesting files, extracting text, generating tags and summaries,
    and storing everything in the database.
    """
    
    # Supported MIME types
    SUPPORTED_TYPES = {
        # Images
        'image/jpeg': 'ocr',
        'image/jpg': 'ocr', 
        'image/png': 'ocr',
        'image/gif': 'ocr',
        'image/bmp': 'ocr',
        'image/tiff': 'ocr',
        'image/webp': 'ocr',
        
        # Documents
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        
        # Text files
        'text/plain': 'text',
        'text/markdown': 'text',
        'text/csv': 'text',
    }
    
    def __init__(self, llm_provider: LLMProvider):
        """
        Initialize the IngestAgent with an LLM provider.
        
        Args:
            llm_provider: LLM provider for generating tags and summaries
        """
        self.llm_provider = llm_provider
        self.blobs_dir = Path("./data/blobs")
        self.blobs_dir.mkdir(parents=True, exist_ok=True)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        if not TESSERACT_AVAILABLE:
            logger.warning("pytesseract not available. OCR functionality will be limited.")
        if not PYMUPDF_AVAILABLE and not PDFMINER_AVAILABLE:
            logger.warning("No PDF libraries available. PDF processing will not work.")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX processing will not work.")
    
    def is_supported(self, mime_type: str) -> bool:
        """Check if the MIME type is supported for text extraction"""
        return mime_type in self.SUPPORTED_TYPES
    
    def ingest_file(
        self, 
        file_data: bytes, 
        filename: str,
        mime_type: str,
        db: Session,
        title: Optional[str] = None
    ) -> Tuple[Optional[Document], List[str]]:
        """
        Ingest a file: extract text, generate tags and summary, store in database.
        
        Args:
            file_data: Raw file data as bytes
            filename: Original filename
            mime_type: MIME type of the file
            db: Database session
            title: Optional custom title for the document
            
        Returns:
            Tuple of (Document object or None if failed, list of error messages)
        """
        errors = []
        
        # Input validation
        if not file_data:
            errors.append("File data cannot be empty")
            return None, errors
        
        if not filename:
            errors.append("Filename cannot be empty")
            return None, errors
        
        if not mime_type:
            errors.append("MIME type cannot be empty")
            return None, errors
        
        try:
            # Get file metadata
            file_size = len(file_data)
            
            # File size validation
            if file_size > MAX_FILE_SIZE:
                errors.append(f"File too large: {file_size} bytes (max: {MAX_FILE_SIZE} bytes)")
                return None, errors
            
            # Calculate content hash for deduplication
            content_hash = compute_bytes_hash(file_data)
            
            # Check if document already exists
            existing_doc = DocumentCRUD.get_by_hash(db, content_hash)
            if existing_doc:
                errors.append(f"Document with this content already exists: {existing_doc.title}")
                return existing_doc, errors
            
            # Extract text from file data
            logger.info(f"Extracting text from {filename}...")
            extracted_text = self._extract_text(file_data, mime_type, filename)
            if not extracted_text:
                logger.warning(f"Could not extract text from {filename}, creating document without content")
                extracted_text = ""  # Create empty text instead of failing
                errors.append(f"Could not extract text from {filename} - document created without content")
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters")
            
            # Generate title if not provided
            if not title:
                title = Path(filename).stem
            
            # Generate summary using LLM
            summary = ""
            if self.llm_provider.is_available():
                try:
                    logger.info("Generating summary using LLM...")
                    # For images with no extracted text, provide context about the image
                    if not extracted_text and mime_type.startswith('image/'):
                        image_context = f"Image file: {filename}, MIME type: {mime_type}, Size: {len(file_data)} bytes. This appears to be an image document that may contain text, documents, or other visual information. Please analyze the image content and provide a summary based on what you can determine from the filename and context."
                        summary = self.llm_provider.summarize(image_context)
                    else:
                        summary = self.llm_provider.summarize(extracted_text)
                    logger.info(f"Generated summary: {summary[:MAX_SUMMARY_PREVIEW]}...")
                except Exception as e:
                    errors.append(f"Failed to generate summary: {str(e)}")
            else:
                logger.warning("LLM not available, skipping summary generation")
                errors.append("OpenAI API key not configured - summary generation skipped")
            
            # Save file to disk
            file_extension = Path(filename).suffix or '.bin'
            blob_filename = f"{content_hash}{file_extension}"
            blob_path = self.blobs_dir / blob_filename
            
            # Write file data to disk
            blob_path.write_bytes(file_data)
            logger.info(f"File saved to: {blob_path}")
            
            # Create document record
            current_time = int(time.time() * 1000)
            document_data = DocumentCreate(
                title=title,
                mime_type=mime_type,
                size_bytes=file_size,
                content_hash=content_hash,
                storage_path=str(blob_path),  # Real file path
                summary=summary if summary else None,
                created_at=current_time,
                imported_at=current_time
            )
            
            # Save document to database
            document = DocumentCRUD.create(db, document_data)
            
            # Generate and assign tags using LLM
            tags = []
            if self.llm_provider.is_available():
                try:
                    logger.info("Generating tags using LLM...")
                    # For images with no extracted text, provide context about the image
                    if not extracted_text and mime_type.startswith('image/'):
                        image_context = f"Image file: {filename}, MIME type: {mime_type}, Size: {len(file_data)} bytes. This appears to be an image document that may contain text, documents, or other visual information. Please analyze the image content and generate relevant tags based on what you can determine from the filename and context."
                        tag_names = self.llm_provider.generate_tags(image_context)
                    else:
                        tag_names = self.llm_provider.generate_tags(extracted_text)
                    
                    if tag_names:
                        logger.info(f"Generated tags: {tag_names}")
                        
                        # Add tags to tags table (get or create)
                        for tag_name in tag_names:
                            TagCRUD.get_or_create(db, tag_name)
                        
                        # Update document with tags as JSON
                        document.tags = json.dumps(tag_names)
                        db.commit()
                        tags = tag_names
                    else:
                        logger.warning("No tags generated by LLM")
                except Exception as e:
                    errors.append(f"Failed to generate tags: {str(e)}")
            else:
                logger.warning("LLM not available, skipping tag generation")
                errors.append("OpenAI API key not configured - tag generation skipped")
            
            return document, errors
            
        except Exception as e:
            errors.append(f"Unexpected error during ingestion: {str(e)}")
            return None, errors
    
    def _extract_text(self, file_data: bytes, mime_type: str, filename: str) -> Optional[str]:
        """
        Extract text from raw file data based on its MIME type
        
        Args:
            file_data: Raw file data as bytes
            mime_type: MIME type of the file
            filename: Original filename for debugging
            
        Returns:
            Extracted text or None if extraction failed
        """
        if not self.is_supported(mime_type):
            logger.warning(f"Unsupported file type: {mime_type}")
            return None
        
        extraction_method = self.SUPPORTED_TYPES[mime_type]
        
        try:
            if extraction_method == 'ocr':
                # Try Vision API first for images, fallback to OCR
                if mime_type.startswith('image/'):
                    extracted_text = self._extract_with_vision_api(file_data, filename)
                    if not extracted_text:
                        logger.warning("Vision API failed, trying OCR fallback...")
                        extracted_text = self._extract_with_ocr(file_data)
                    return extracted_text
                else:
                    return self._extract_with_ocr(file_data)
            elif extraction_method == 'pdf':
                return self._extract_from_pdf(file_data)
            elif extraction_method == 'docx':
                return self._extract_from_docx(file_data)
            elif extraction_method == 'text':
                return self._extract_from_text(file_data)
            else:
                logger.error(f"Unknown extraction method: {extraction_method}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return None
    
    def _extract_with_vision_api(self, file_data: bytes, filename: str) -> Optional[str]:
        """Extract text from image using OpenAI Vision API"""
        if not self.llm_provider.is_available():
            logger.warning("Vision API not available - OpenAI API key not configured")
            return None
        
        try:
            logger.info(f"Using Vision API for image: {filename}")
            extracted_text = self.llm_provider.extract_text_from_image(file_data, filename)
            
            if not extracted_text:
                logger.warning("No text extracted by Vision API")
                return None
            
            logger.info(f"Vision API extracted {len(extracted_text)} characters from image")
            logger.debug(f"Vision API text preview: {extracted_text[:MAX_TEXT_PREVIEW]}...")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Vision API extraction failed: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return None

    def _extract_with_ocr(self, file_data: bytes) -> Optional[str]:
        """Extract text from image bytes using OCR (fallback)"""
        if not TESSERACT_AVAILABLE:
            logger.warning("OCR not available - pytesseract not installed")
            return None
        
        try:
            # Open image from bytes
            image = Image.open(BytesIO(file_data))
            logger.debug(f"Image format: {image.format}, mode: {image.mode}, size: {image.size}")
            
            # Handle MPO and other problematic formats
            if image.format == 'MPO':
                logger.info("Converting MPO to JPEG format for OCR")
                # Convert MPO to RGB and then save as JPEG in memory
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Save as JPEG in memory
                import io
                jpeg_buffer = io.BytesIO()
                image.save(jpeg_buffer, format='JPEG', quality=95)
                jpeg_buffer.seek(0)
                image = Image.open(jpeg_buffer)
                logger.debug(f"Converted to JPEG: {image.format}, mode: {image.mode}")
            
            # Convert to RGB if necessary (for some image formats)
            if image.mode not in ['RGB', 'L']:
                logger.debug(f"Converting image from {image.mode} to RGB")
                image = image.convert('RGB')
            
            # Apply generic image preprocessing to improve OCR
            logger.debug("Applying image preprocessing for better OCR...")
            try:
                # Convert to grayscale for better contrast
                if image.mode != 'L':
                    gray_image = image.convert('L')
                else:
                    gray_image = image
                
                # Apply contrast enhancement for better text recognition
                from PIL import ImageEnhance
                enhancer = ImageEnhance.Contrast(gray_image)
                enhanced_image = enhancer.enhance(1.5)  # Moderate contrast increase
                
                # Try with enhanced image first
                image = enhanced_image
                logger.debug("Applied contrast enhancement for better OCR")
            except Exception as e:
                logger.debug(f"Image preprocessing failed: {e}, using original image")
            
            # Try different OCR configurations for better results
            try:
                # Try multiple OCR configurations for better text extraction
                ocr_configs = [
                    '--psm 6',  # Uniform block of text
                    '--psm 4',  # Single column of text
                    '--psm 3',  # Fully automatic page segmentation
                    '--psm 8',  # Single word
                    '--psm 7',  # Single text line
                    ''  # Default
                ]
                
                text = ""
                for config in ocr_configs:
                    try:
                        if config:
                            text = pytesseract.image_to_string(image, lang='eng', config=config)
                        else:
                            text = pytesseract.image_to_string(image, lang='eng')
                        
                        if text and len(text.strip()) > 50:
                            logger.debug(f"OCR succeeded with config: {config or 'default'}")
                            break
                    except Exception as e:
                        logger.debug(f"OCR failed with config {config or 'default'}: {e}")
                        continue
                    
            except Exception as ocr_error:
                logger.warning(f"All OCR attempts failed: {ocr_error}")
                text = ""
            
            # Clean up the text
            text = text.strip()
            
            if not text:
                logger.warning(f"No text found in image")
                return None
            
            logger.info(f"OCR extracted {len(text)} characters from image")
            logger.debug(f"OCR text preview: {text[:MAX_TEXT_PREVIEW]}...")
            return text
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return None
    
    def _extract_from_pdf(self, file_data: bytes) -> Optional[str]:
        """Extract text from PDF bytes"""
        # Try PyMuPDF first (faster)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(stream=file_data, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                if text.strip():
                    logger.info(f"PyMuPDF extracted {len(text)} characters from PDF")
                    return text.strip()
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Fallback to pdfminer
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract(BytesIO(file_data))
                if text.strip():
                    logger.info(f"pdfminer extracted {len(text)} characters from PDF")
                    return text.strip()
            except Exception as e:
                logger.warning(f"pdfminer extraction failed: {e}")
        
        logger.warning(f"No PDF extraction method available")
        return None
    
    def _extract_from_docx(self, file_data: bytes) -> Optional[str]:
        """Extract text from DOCX bytes"""
        if not DOCX_AVAILABLE:
            logger.warning("DOCX extraction not available - python-docx not installed")
            return None
        
        try:
            doc = DocxDocument(BytesIO(file_data))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            text = text.strip()
            if text:
                logger.info(f"DOCX extracted {len(text)} characters")
                return text
            else:
                logger.warning(f"No text found in DOCX")
                return None
                
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None
    
    def _extract_from_text(self, file_data: bytes) -> Optional[str]:
        """Extract text from raw text bytes"""
        try:
            # Try different encodings
            encodings = TEXT_ENCODINGS
            
            for encoding in encodings:
                try:
                    text = file_data.decode(encoding)
                    
                    if text.strip():
                        logger.info(f"Text extracted {len(text)} characters using {encoding}")
                        return text.strip()
                        
                except UnicodeDecodeError:
                    continue
            
            logger.warning(f"Could not decode text data")
            return None
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return None
    
    
