"""
Text extraction utilities for various file formats including OCR for images
"""
import mimetypes
from pathlib import Path
from typing import Optional

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import PyMuPDF as fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TextExtractor:
    """
    Handles text extraction from various file formats including OCR for images
    """
    
    # Supported MIME types
    SUPPORTED_TYPES = {
        # Images
        'image/jpeg': 'ocr',
        'image/jpg': 'ocr', 
        'image/png': 'ocr',
        'image/gif': 'ocr',
        'image/bmp': 'ocr',
        'image/tiff': 'ocr',
        'image/webp': 'ocr',
        
        # Documents
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        
        # Text files
        'text/plain': 'text',
        'text/markdown': 'text',
        'text/csv': 'text',
    }
    
    def __init__(self):
        """Initialize the text extractor"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        if not TESSERACT_AVAILABLE:
            print("Warning: pytesseract not available. OCR functionality will be limited.")
        if not PYMUPDF_AVAILABLE and not PDFMINER_AVAILABLE:
            print("Warning: No PDF libraries available. PDF processing will not work.")
        if not DOCX_AVAILABLE:
            print("Warning: python-docx not available. DOCX processing will not work.")
    
    def is_supported(self, mime_type: str) -> bool:
        """Check if the MIME type is supported for text extraction"""
        return mime_type in self.SUPPORTED_TYPES
    
    def extract_text(self, file_data: bytes, mime_type: str, filename: str = "unknown") -> Optional[str]:
        """
        Extract text from raw file data based on its MIME type
        
        Args:
            file_data: Raw file data as bytes
            mime_type: MIME type of the file
            filename: Original filename for debugging
            
        Returns:
            Extracted text or None if extraction failed
        """
        if not self.is_supported(mime_type):
            print(f"Unsupported file type: {mime_type}")
            return None
        
        extraction_method = self.SUPPORTED_TYPES[mime_type]
        
        try:
            if extraction_method == 'ocr':
                return self._extract_with_ocr_from_bytes(file_data)
            elif extraction_method == 'pdf':
                return self._extract_from_pdf_from_bytes(file_data)
            elif extraction_method == 'docx':
                return self._extract_from_docx_from_bytes(file_data)
            elif extraction_method == 'text':
                return self._extract_from_text_from_bytes(file_data)
            else:
                print(f"Unknown extraction method: {extraction_method}")
                return None
                
        except Exception as e:
            print(f"Error extracting text from {filename}: {e}")
            return None
    
    def _extract_with_ocr(self, file_path: Path) -> Optional[str]:
        """Extract text from images using OCR"""
        if not TESSERACT_AVAILABLE:
            print("OCR not available - pytesseract not installed")
            return None
        
        try:
            # Open image with PIL
            image = Image.open(file_path)
            
            # Convert to RGB if necessary (for some image formats)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image, lang='eng')
            
            # Clean up the text
            text = text.strip()
            
            if not text:
                print(f"No text found in image: {file_path}")
                return None
            
            print(f"OCR extracted {len(text)} characters from {file_path}")
            return text
            
        except Exception as e:
            print(f"OCR extraction failed for {file_path}: {e}")
            return None
    
    def _extract_from_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF files"""
        # Try PyMuPDF first (faster)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                if text.strip():
                    print(f"PyMuPDF extracted {len(text)} characters from {file_path}")
                    return text.strip()
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}")
        
        # Fallback to pdfminer
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract(str(file_path))
                if text.strip():
                    print(f"pdfminer extracted {len(text)} characters from {file_path}")
                    return text.strip()
            except Exception as e:
                print(f"pdfminer extraction failed: {e}")
        
        print(f"No PDF extraction method available for {file_path}")
        return None
    
    def _extract_from_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            print("DOCX extraction not available - python-docx not installed")
            return None
        
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            text = text.strip()
            if text:
                print(f"DOCX extracted {len(text)} characters from {file_path}")
                return text
            else:
                print(f"No text found in DOCX: {file_path}")
                return None
                
        except Exception as e:
            print(f"DOCX extraction failed for {file_path}: {e}")
            return None
    
    def _extract_from_text(self, file_path: Path) -> Optional[str]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    if text.strip():
                        print(f"Text extracted {len(text)} characters from {file_path} using {encoding}")
                        return text.strip()
                        
                except UnicodeDecodeError:
                    continue
            
            print(f"Could not decode text file: {file_path}")
            return None
            
        except Exception as e:
            print(f"Text extraction failed for {file_path}: {e}")
            return None
    
    def _extract_with_ocr_from_bytes(self, file_data: bytes) -> Optional[str]:
        """Extract text from image bytes using OCR"""
        if not TESSERACT_AVAILABLE:
            print("OCR not available - pytesseract not installed")
            return None
        
        try:
            from io import BytesIO
            # Open image from bytes
            image = Image.open(BytesIO(file_data))
            
            # Convert to RGB if necessary (for some image formats)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image, lang='eng')
            
            # Clean up the text
            text = text.strip()
            
            if not text:
                print(f"No text found in image")
                return None
            
            print(f"OCR extracted {len(text)} characters from image")
            return text
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return None
    
    def _extract_from_pdf_from_bytes(self, file_data: bytes) -> Optional[str]:
        """Extract text from PDF bytes"""
        # Try PyMuPDF first (faster)
        if PYMUPDF_AVAILABLE:
            try:
                from io import BytesIO
                doc = fitz.open(stream=file_data, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                if text.strip():
                    print(f"PyMuPDF extracted {len(text)} characters from PDF")
                    return text.strip()
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}")
        
        # Fallback to pdfminer
        if PDFMINER_AVAILABLE:
            try:
                from io import BytesIO
                text = pdfminer_extract(BytesIO(file_data))
                if text.strip():
                    print(f"pdfminer extracted {len(text)} characters from PDF")
                    return text.strip()
            except Exception as e:
                print(f"pdfminer extraction failed: {e}")
        
        print(f"No PDF extraction method available")
        return None
    
    def _extract_from_docx_from_bytes(self, file_data: bytes) -> Optional[str]:
        """Extract text from DOCX bytes"""
        if not DOCX_AVAILABLE:
            print("DOCX extraction not available - python-docx not installed")
            return None
        
        try:
            from io import BytesIO
            doc = DocxDocument(BytesIO(file_data))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            text = text.strip()
            if text:
                print(f"DOCX extracted {len(text)} characters")
                return text
            else:
                print(f"No text found in DOCX")
                return None
                
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return None
    
    def _extract_from_text_from_bytes(self, file_data: bytes) -> Optional[str]:
        """Extract text from raw text bytes"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_data.decode(encoding)
                    
                    if text.strip():
                        print(f"Text extracted {len(text)} characters using {encoding}")
                        return text.strip()
                        
                except UnicodeDecodeError:
                    continue
            
            print(f"Could not decode text data")
            return None
            
        except Exception as e:
            print(f"Text extraction failed: {e}")
            return None
